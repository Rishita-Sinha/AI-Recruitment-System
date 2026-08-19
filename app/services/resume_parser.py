import pdfplumber
import numpy as np

from docx import Document
from rapidocr_onnxruntime import RapidOCR


def clean_text(text: str) -> str:
    """
    Clean extracted resume text by removing hidden characters
    that can cause PostgreSQL insertion errors.
    """
    if not text:
        return ""

    return (
        text.replace("\x00", "")
            .replace("\ufeff", "")
            .replace("\u200b", "")
            .replace("\r", "\n")
            .strip()
    )


def extract_text_with_ocr(pdf_path: str) -> str:
    """
    Extract text from scanned/image PDFs using RapidOCR.
    """
    ocr = RapidOCR()
    all_text = ""

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            try:
                # Render page as high-resolution image
                image = page.to_image(resolution=300).original

                # Convert PIL Image to NumPy array
                image = np.array(image)

                result, _ = ocr(image)

                if result:
                    page_text = "\n".join([line[1] for line in result])
                    all_text += page_text + "\n"

            except Exception as e:
                print(f"OCR failed on page: {e}")

    return clean_text(all_text)


def extract_text_from_pdf(pdf_path: str) -> str:
    """
    First try normal text extraction.
    If nothing is found, automatically fall back to OCR.
    """
    extracted_text = ""

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()

            if page_text:
                extracted_text += page_text + "\n"

    extracted_text = clean_text(extracted_text)

    # OCR fallback
    if not extracted_text.strip():
        print("No selectable text found. Using OCR...")
        extracted_text = extract_text_with_ocr(pdf_path)

    return extracted_text


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from DOCX files including:
    - Normal paragraphs
    - Tables
    - Headers
    - Footers
    - Text boxes / shapes
    """

    doc = Document(docx_path)

    extracted_parts = []

    # ----------------------------------------
    # Normal paragraphs
    # ----------------------------------------
    for para in doc.paragraphs:
        text = para.text.strip()

        if text:
            extracted_parts.append(text)

    # ----------------------------------------
    # Tables
    # ----------------------------------------
    for table in doc.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                cell_text = cell.text.strip()

                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                extracted_parts.append(
                    " | ".join(row_text)
                )

    # ----------------------------------------
    # Text boxes / shapes
    # ----------------------------------------
    for element in doc.element.body.iter():

        if element.tag.endswith("}t"):
            text = element.text

            if text and text.strip():
                extracted_parts.append(
                    text.strip()
                )

    # ----------------------------------------
    # Headers
    # ----------------------------------------
    for section in doc.sections:

        for para in section.header.paragraphs:

            text = para.text.strip()

            if text:
                extracted_parts.append(text)

    # ----------------------------------------
    # Footers
    # ----------------------------------------
    for section in doc.sections:

        for para in section.footer.paragraphs:

            text = para.text.strip()

            if text:
                extracted_parts.append(text)

    return clean_text(
        "\n".join(extracted_parts)
    )


def extract_resume_text(file_path: str) -> str:
    """
    Detect file type and extract text accordingly.
    """
    if file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)

    elif file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)

    else:
        raise ValueError(f"Unsupported file format: {file_path}")